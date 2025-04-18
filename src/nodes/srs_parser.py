from typing import Dict, Any, Annotated, Callable, Tuple, List
from langgraph.graph import MessageGraph, StateGraph
from groq import Groq
import os
from docx import Document
import json
import re
from langsmith import Client

# Constants
MODEL_NAME = "mistral-saba-24b"  # Using latest Mixtral model on Groq

def get_groq_client():
    """Get or initialize Groq client."""
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise ValueError("GROQ_API_KEY environment variable must be set")
    return Groq(api_key=api_key)

def process_docx(file_path: str) -> str:
    """Extract text content from .docx file."""
    doc = Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

async def analyze_requirements(content: str) -> dict:
    """Extract functional requirements from SRS content."""
    prompt = f"""You are a software engineer analyzing a Software Requirements Specification (SRS) document.
    Your task is to extract structured information and return it in valid JSON format.

    Format your response exactly like this example:
    {{
        "functional_requirements": [
            "User registration",
            "Password reset functionality"
        ],
        "api_endpoints": [
            {{
                "path": "/api/users",
                "method": "POST",
                "description": "Create new user"
            }}
        ],
        "db_schema": {{
            "tables": [
                {{
                    "name": "users",
                    "fields": ["id", "username", "email"]
                }}
            ]
        }},
        "auth_requirements": {{
            "type": "JWT",
            "features": ["RBAC"]
        }}
    }}

    Now analyze this SRS document and return the information in the same JSON structure:
    {content}

    Remember: The response must be valid JSON, use double quotes for strings, and follow the exact structure shown above."""
    
    response = get_groq_client().chat.completions.create(
        model=MODEL_NAME,
        messages=[{
            "role": "user",
            "content": prompt
        }],
        temperature=0.2,
        max_tokens=4000,
    )
    
    raw_response = response.choices[0].message.content
    
    # Extract JSON from the response - handles cases where LLM adds extra text
    try:
        # Find JSON content (typically between curly braces)
        json_start = raw_response.find('{')
        json_end = raw_response.rfind('}') + 1
        if json_start >= 0 and json_end > json_start:
            json_content = raw_response[json_start:json_end]
            return json.loads(json_content)
        else:
            # If no JSON format detected, try parsing the whole response
            return json.loads(raw_response)
    except json.JSONDecodeError:
        # Fallback structure if JSON parsing fails
        return {
            "functional_requirements": ["Failed to parse requirements"],
            "api_endpoints": [],
            "db_schema": {"tables": []},
            "auth_requirements": {"type": "Unknown", "features": []}
        }

async def srs_parser(state: Dict[str, Any]) -> Dict[str, Any]:
    """Parse SRS document and extract requirements."""
    try:
        if not state["srs_content"]:
            raise ValueError("Empty SRS content")

        if state["srs_content"].endswith(".docx"):
            content = process_docx(state["srs_content"])
        else:
            content = state["srs_content"]

        requirements = await analyze_requirements(content)
        
        # Basic validation of the requirements structure
        required_keys = ["functional_requirements", "api_endpoints", "db_schema", "auth_requirements"]
        missing_keys = [key for key in required_keys if key not in requirements]
        
        if missing_keys:
            for key in missing_keys:
                # Add missing keys with default values
                if key == "functional_requirements":
                    requirements[key] = []
                elif key == "api_endpoints":
                    requirements[key] = []
                elif key == "db_schema":
                    requirements[key] = {"tables": []}
                elif key == "auth_requirements":
                    requirements[key] = {"type": "Unknown", "features": []}
            
            state["logs"].append(f"Warning: Added missing keys in requirements: {', '.join(missing_keys)}")
        
        # Update state with extracted requirements
        state["requirements"] = requirements
        state["logs"].append("Successfully parsed SRS document")
        
        return state
        
    except Exception as e:
        state["errors"].append(f"SRS parsing error: {str(e)}")
        state["logs"].append(f"Error in SRS parsing: {str(e)}")
        return state

class SRSParserNode:
    """Node for parsing Software Requirements Specification documents."""

    def __init__(self):
        self.langsmith_client = Client()

    async def extract_requirements(self, doc_path: str) -> Dict[str, Any]:
        """Extract requirements from a Word document."""
        try:
            doc = docx.Document(doc_path)
            
            # Initialize requirements dictionary
            requirements = {
                "endpoints": [],
                "models": [],
                "auth": {},
                "features": []
            }
            
            # Extract API endpoints using regex patterns
            endpoint_pattern = re.compile(r'/api/\w+(?:/\w+)*')
            
            # Extract data models - typically CamelCase words followed by "model"
            model_pattern = re.compile(r'([A-Z][a-z0-9]+)+(?:\s+[Mm]odel)')
            
            # Extract features - use bullet points or numbered lists
            feature_pattern = re.compile(r'(?:•|\d+\.)\s+(.+?)(?=(?:•|\d+\.)|$)', re.DOTALL)
            
            # Process each paragraph in the document
            for para in doc.paragraphs:
                text = para.text.strip()
                
                # Extract endpoints
                for endpoint in endpoint_pattern.findall(text):
                    if endpoint not in requirements["endpoints"]:
                        requirements["endpoints"].append(endpoint)
                
                # Extract models
                for model_match in model_pattern.finditer(text):
                    model_name = model_match.group().split()[0]  # Get just the model name
                    if model_name not in requirements["models"]:
                        requirements["models"].append(model_name)
                
                # Extract features
                for feature_match in feature_pattern.finditer(text):
                    feature = feature_match.group(1).strip()
                    if feature and feature not in requirements["features"]:
                        requirements["features"].append(feature)
                
                # Check for authentication information
                if "authentication" in text.lower() or "auth" in text.lower():
                    # Very basic extraction
                    if "jwt" in text.lower() or "token" in text.lower():
                        requirements["auth"]["type"] = "JWT"
                    elif "oauth" in text.lower():
                        requirements["auth"]["type"] = "OAuth"
                    elif "basic" in text.lower():
                        requirements["auth"]["type"] = "Basic"
                    
                    # Try to extract expiry time if mentioned
                    expiry_match = re.search(r'(\d+)\s*(?:second|minute|hour|day)', text.lower())
                    if expiry_match:
                        time_value = int(expiry_match.group(1))
                        time_unit = expiry_match.group(2)
                        
                        # Convert to seconds
                        if "minute" in time_unit:
                            time_value *= 60
                        elif "hour" in time_unit:
                            time_value *= 3600
                        elif "day" in time_unit:
                            time_value *= 86400
                            
                        requirements["auth"]["expiry"] = time_value
            
            return requirements
            
        except Exception as e:
            raise Exception(f"Error extracting requirements: {e}")

    async def run(self, state: Dict[str, Any]) -> Tuple[Dict[str, Any], str]:
        """Run the SRS parser node."""
        try:
            # Create a LangSmith run
            run = self.langsmith_client.create_run(
                "srs_parser_run",
                inputs={"srs_path": state.get("srs_path")}
            )
            
            # Extract requirements from the document
            requirements = await self.extract_requirements(state.get("srs_path"))
            
            # Update state with extracted requirements
            state["requirements"] = requirements
            state["logs"].append("Successfully parsed SRS document and extracted requirements")
            
            # Update LangSmith run with success status
            self.langsmith_client.update_run(
                run.id,
                outputs={"status": "success", "requirements": requirements}
            )
            
            # Return updated state and next node
            return state, "test_generator"
            
        except Exception as e:
            # Handle errors
            error_msg = f"Error parsing SRS document: {e}"
            state["errors"].append(error_msg)
            state["logs"].append("Failed to parse SRS document")
            
            # Log error to LangSmith
            self.langsmith_client.update_run(
                run_id="srs_parser_run",
                error=str(e)
            )
            
            # Return error state
            return state, "error_handler"
